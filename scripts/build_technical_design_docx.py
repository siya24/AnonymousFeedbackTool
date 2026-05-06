from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "TECHNICAL_DESIGN.md"
OUT = ROOT / "TECHNICAL_DESIGN.docx"


def clean_inline(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.strip()


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [clean_inline(c) for c in stripped.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return

    header = rows[0]
    data_rows = [r for r in rows[2:] if any(c.strip("-: ") for c in r)]

    col_count = max(len(header), *(len(r) for r in data_rows), 1)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = "Table Grid"

    for i, val in enumerate(header):
        table.rows[0].cells[i].text = val

    for r_idx, row in enumerate(data_rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build_doc(md_text: str) -> Document:
    doc = Document()

    base_style = doc.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    title = doc.add_paragraph()
    title_run = title.add_run("Anonymous Feedback Tool - Technical Design")
    title_run.bold = True
    title_run.font.size = Pt(18)

    meta = doc.add_paragraph(
        f"Generated from TECHNICAL_DESIGN.md on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.name = "Calibri"

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line.rstrip("\n"))
            i += 1
            continue

        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        if line.strip() in {"---", "***", "___"}:
            i += 1
            continue

        if line.lstrip().startswith("|"):
            table_block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_block.append(lines[i])
                i += 1
            add_table(doc, table_block)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_inline(heading_match.group(2))
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", line.strip())
        if ordered_match:
            doc.add_paragraph(clean_inline(ordered_match.group(1)), style="List Number")
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line.strip())
        if bullet_match:
            doc.add_paragraph(clean_inline(bullet_match.group(1)), style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(clean_inline(line))
        i += 1

    return doc


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Source markdown not found: {SRC}")

    md_text = SRC.read_text(encoding="utf-8", errors="replace")
    doc = build_doc(md_text)
    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
